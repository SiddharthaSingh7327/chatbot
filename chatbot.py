import os
import sys
from pathlib import Path
import google.generativeai as genai # type: ignore
from typing import Optional
import json
import PyPDF2 # type: ignore
from docx import Document as DocxDocument # type: ignore
import io
from PIL import Image # type:ignore
import fitz  #type:ignore
from fast_graphrag import GraphRAG # type:ignore
import shutil
class EnhancedDocumentQAChatbot:
    def __init__(self, api_key: str):
        """Initialize the chatbot with Gemini API key."""
        genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel('gemini-2.5-pro')
        self.document_content = ""
        self.document_name = ""
        self.document_images = []  
        self.chat_history = []
        self.graph_rag_instance = None
        self.working_dir = None
    def load_document(self, file_path: str) -> bool:
        """Load a document from file path and initialize GraphRAG."""
        try:
            path = Path(file_path)
            if not path.exists():
                print(f" File not found: {file_path}")
                return False
            self.document_content = ""
            self.document_images = []
            if self.working_dir and os.path.exists(self.working_dir):
                shutil.rmtree(self.working_dir)
            self.working_dir = f"./graphrag_data_{path.stem}"
            if path.suffix.lower() in ['.txt', '.md', '.py', '.js', '.html', '.css', '.json']:
                with open(path, 'r', encoding='utf-8') as file:
                    self.document_content = file.read()
            elif path.suffix.lower() == '.pdf':
                text_content, images = self._extract_pdf_content_with_images(path)
                self.document_content = text_content
                self.document_images = images
                if not self.document_content and not self.document_images:
                    print(" Could not extract content from PDF file.")
                    return False
            elif path.suffix.lower() in ['.docx', '.doc']:
                if path.suffix.lower() == '.docx':
                    self.document_content = self._extract_docx_text(path)
                else:
                    print("  .doc files are not supported. Please convert to .docx format.")
                    return False
                if not self.document_content:
                    print("  Could not extract text from Word document.")
                    return False
            elif path.suffix.lower() in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff']:
                image = Image.open(path)
                self.document_images = [image]
                self.document_content = f"Image file: {path.name}"
            else:
                try:
                    with open(path, 'r', encoding='utf-8') as file:
                        self.document_content = file.read()
                except UnicodeDecodeError:
                    print(f"  Cannot read file {file_path}. Unsupported file format.")
                    return False
            self.document_name = path.name
            print(f"  Document '{self.document_name}' loaded successfully!")
            print(f"  Text content: {len(self.document_content)} characters")
            print(f"  Images extracted: {len(self.document_images)}")
            if self.document_content:
                self._initialize_and_build_graphrag()
            return True
        except Exception as e:
            print(f"  Error loading document: {e}")
            return False
    def _initialize_and_build_graphrag(self):
        """Initializes GraphRAG and builds the knowledge graph."""
        print("\nInitializing and building knowledge graph with fast-graphrag...")
        try:
            DOMAIN = "This document contains various information. Extract key entities and their relationships."
            EXAMPLE_QUERIES = [
                "Summarize the main points of the document.",
                "What are the key entities mentioned?",
                "Describe the relationships between the main entities."
            ]
            ENTITY_TYPES = ["Person", "Organization", "Location", "Date", "Topic", "Concept"]
            self.graph_rag_instance = GraphRAG(
                working_dir=self.working_dir,
                domain=DOMAIN,
                example_queries="\n".join(EXAMPLE_QUERIES),
                entity_types=ENTITY_TYPES
            )
            self.graph_rag_instance.insert(self.document_content)
            print("  Knowledge graph built successfully!")
        except Exception as e:
            print(f"  Error initializing or building GraphRAG: {e}")
            self.graph_rag_instance = None
    def regenerate_graph(self):
        """Forces regeneration of the knowledge graph."""
        if not self.document_content:
            print("  No document loaded to regenerate the graph for.")
            return
        if self.working_dir and os.path.exists(self.working_dir):
            shutil.rmtree(self.working_dir)
        self._initialize_and_build_graphrag()
    def _extract_pdf_content_with_images(self, pdf_path: Path) -> tuple:
        """Extract both text and images from PDF using PyMuPDF."""
        try:
            text_content = ""
            images = []
            pdf_document = fitz.open(str(pdf_path))
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                page_text = page.get_text()
                if page_text.strip():
                    text_content += f"\n--- Page {page_num + 1} ---\n"
                    text_content += page_text
                image_list = page.get_images()
                for img_index, img in enumerate(image_list):
                    try:
                        xref = img[0]
                        pix = fitz.Pixmap(pdf_document, xref)
                        if pix.n - pix.alpha < 4:
                            img_data = pix.tobytes("png")
                            pil_image = Image.open(io.BytesIO(img_data))
                            images.append(pil_image)
                            print(f"  Extracted image {len(images)} from page {page_num + 1}")
                        pix = None
                    except Exception as e:
                        print(f"⚠️ Could not extract image {img_index} from page {page_num + 1}: {e}")
            pdf_document.close()
            return text_content.strip(), images
        except Exception as e:
            print(f"  Error reading PDF with PyMuPDF: {e}")
            fallback_text = self._extract_pdf_text(pdf_path)
            return fallback_text, []
    def _extract_pdf_text(self, pdf_path: Path) -> str:
        """Fallback method to extract text from PDF using PyPDF2."""
        try:
            text = ""
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += f"\n--- Page {page_num + 1} ---\n"
                            text += page_text
                    except Exception as e:
                        print(f"  Warning: Could not extract text from page {page_num + 1}: {e}")
                if not text.strip():
                    print("  Warning: No text could be extracted from the PDF.")
                    return ""
            return text.strip()
        except Exception as e:
            print(f"  Error reading PDF: {e}")
            return ""
    def _extract_docx_text(self, docx_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            doc = DocxDocument(docx_path)
            text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            return text.strip()
        except Exception as e:
            print(f"  Error reading DOCX: {e}")
            return ""
    def ask_question(self, question: str) -> str:
        """Ask a question using the appropriate engine (GraphRAG or Gemini)."""
        if not self.document_content and not self.document_images:
            return "  No document loaded. Please load a document first."
        if self.graph_rag_instance and not self.document_images:
            print("  Using fast-graphrag for a text-based query...")
            try:
                response = self.graph_rag_instance.query(question)
                answer = response.response
                self.chat_history.append({"question": question, "answer": answer})
                return answer
            except Exception as e:
                return f"  Error querying with GraphRAG: {e}"
        print("  Using Gemini API for multimodal query...")
        try:
            content_parts = []
            if self.document_content:
                text_prompt = f"""
Document Name: {self.document_name}
Document Content:
---
{self.document_content}
---
"""
                content_parts.append(text_prompt)
            if self.document_images:
                content_parts.append("The document also contains the following images:")
                for i, image in enumerate(self.document_images):
                    content_parts.append(f"Image {i+1}:")
                    content_parts.append(image)
            if self.chat_history:
                content_parts.append(f"\nPrevious conversation:\n{self._format_chat_history()}")
            question_prompt = f"""
Question: {question}
Instructions:
- Analyze both text content AND any visual elements (images, charts, etc.).
- Provide a detailed and accurate answer based on ALL available content.
"""
            content_parts.append(question_prompt)
            response = self.model.generate_content(content_parts)
            answer = response.text
            self.chat_history.append({"question": question, "answer": answer})
            return answer
        except Exception as e:
            return f"  Error generating response with Gemini: {e}"
    def _format_chat_history(self) -> str:
        """Format chat history for context."""
        if not self.chat_history:
            return "No previous conversation."
        history = ""
        for i, entry in enumerate(self.chat_history[-3:], 1):
            history += f"\nQ{i}: {entry['question']}\nA{i}: {entry['answer']}\n"
        return history
    def clear_history(self):
        """Clear chat history."""
        self.chat_history = []
        print("  Chat history cleared!")
    def show_document_info(self):
        """Show information about the loaded document."""
        if not self.document_content and not self.document_images:
            print("  No document loaded.")
            return
        print(f"\n  Document Information:")
        print(f"   Name: {self.document_name}")
        if self.document_content:
            print(f"   Text Size: {len(self.document_content)} characters")
            preview = self.document_content[:200]
            if len(self.document_content) > 200:
                preview += "..."
            print(f"   Text Preview: {preview}")
        if self.document_images:
            print(f"   Images: {len(self.document_images)} extracted")
            for i, img in enumerate(self.document_images):
                print(f"     Image {i+1}: {img.size[0]}x{img.size[1]} pixels")
def print_welcome():
    """Print welcome message."""
    print("=" * 60)
    print("  ENHANCED DOCUMENT Q&A CHATBOT with fast-graphrag & Gemini")
    print("=" * 60)
    print("This chatbot uses fast-graphrag for text analysis and Gemini for multimodal queries.")
    print("\n  Supported file formats:")
    print("   • Text files: .txt, .md, .py, .js, .html, .css, .json")
    print("   • PDF files: .pdf")
    print("   • Word documents: .docx")
    print("   • Image files: .jpg, .jpeg, .png")
    print("\nCommands:")
    print("     /load <file_path>  - Load a document")
    print("     /info             - Show document information")
    print("     /regen_graph      - Regenerate the knowledge graph")
    print("     /clear            - Clear chat history")
    print("     /help             - Show this help")
    print("     /quit or /exit    - Exit the program")
    print("\nJust type your question to ask about the loaded document!")
    print("=" * 60)
def main():
    api_key = os.environ.get("GEMINI_API_KEY")
    openai_api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        print("  Gemini API key not found as an environment variable.")
        api_key = input("   Please enter your Gemini API key: ").strip()
        if not api_key:
            print("  No Gemini API key provided. Exiting.")
            sys.exit(1)
    if not openai_api_key:
        print("  OpenAI API key not found as an environment variable (required by fast-graphrag).")
        openai_api_key = input("   Please enter your OpenAI API key: ").strip()
        if not openai_api_key:
            print("  No OpenAI API key provided. fast-graphrag may not function correctly.")
    os.environ["OPENAI_API_KEY"] = openai_api_key
    try:
        chatbot = EnhancedDocumentQAChatbot(api_key)
        print("  Gemini API initialized successfully!")
    except Exception as e:
        print(f"  Error initializing Gemini API: {e}")
        sys.exit(1)
    print_welcome()
    while True:
        try:
            user_input = input("\n  Ask me anything: ").strip()
            if not user_input:
                continue
            if user_input.startswith('/'):
                command_parts = user_input.split(' ', 1)
                command = command_parts[0].lower()
                if command in ['/quit', '/exit']:
                    print("  Goodbye!")
                    if chatbot.working_dir and os.path.exists(chatbot.working_dir):
                        shutil.rmtree(chatbot.working_dir)
                    break
                elif command == '/help':
                    print_welcome()
                elif command == '/load':
                    if len(command_parts) < 2:
                        print("  Please provide a file path: /load <file_path>")
                    else:
                        file_path = command_parts[1]
                        chatbot.load_document(file_path)
                elif command == '/info':
                    chatbot.show_document_info()
                elif command == '/clear':
                    chatbot.clear_history()
                elif command == '/regen_graph':
                    chatbot.regenerate_graph()
                else:
                    print("  Unknown command. Type /help for available commands.")
            else:
                if not chatbot.document_content and not chatbot.document_images:
                    print("  Please load a document first using: /load <file_path>")
                    continue
                print("\n  Thinking...")
                answer = chatbot.ask_question(user_input)
                print(f"\n  Answer:\n{answer}")
        except KeyboardInterrupt:
            print("\n\n  Goodbye!")
            if chatbot.working_dir and os.path.exists(chatbot.working_dir):
                shutil.rmtree(chatbot.working_dir)
            break
        except Exception as e:
            print(f"  An unexpected error occurred: {e}")
if __name__=="__main__":
    main()